import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import datetime
import webbrowser
import os
import six
from docx import Document
from docx.shared import Inches
plt.style.use('ggplot')

#convert excel files to csv
def SJM_toCSV():
    pd.read_excel('SJM.xlsx', 'Page 1', index_col=None).to_csv('SJM.csv', encoding='utf-8')
def SJMCallStats_toCSV():
    pd.read_excel('SJ Milli - Previous Week.xls', 'Sheet0', index_col=None).to_csv('SJ Milli - Previous Week.csv', encoding='utf-8')

def convert_toCSV():
    SJM_toCSV()
    SJMCallStats_toCSV()
    if os.path.isfile('SJM.csv'):
        print("The file, SJM.csv, generated successfully")
    else:
        print("The file, SJM.csv, DID NOT generate successfully")
    if os.path.isfile('SJ Milli - Previous Week.csv'):
        print("The file, SJ Milli - Previous Week.csv, generated successfully")
    else:
        print("The file, SJ Milli - Previous Week.csv, DID NOT generate successfully")

#find total calls in call statistics report
def find_total_calls(df):
    totalRows = (len(df.index - 30))
    return (totalRows - 29)

#clean list
def clean_list(df, totalCalls):
    df_toList = df.iloc[-5].tolist()
    del df_toList[0]
    clean1_list = [x for x in df_toList if str(x) != 'nan']
    for x in range(len(clean1_list)):
        for char in "Avg: ": clean1_list[x] = clean1_list[x].replace(char,"")
    clean2_list = [x for x in clean1_list if len(x) < 5 and len(x) > 0]
    clean3_list = [x for x in clean2_list if str(x) != '  ']
    clean3_list.append(totalCalls)
    return list(map(int, clean3_list))

#clean list for averages
def clean_list_averages(df):
    df_list = df.iloc[-5].tolist()
    del df_list[0]
    clean1_list = [x for x in df_list if str(x) != 'nan']
    for x in range(len(clean1_list)):
        for char in "Avg: ": clean1_list[x] = clean1_list[x].replace(char,"")
    clean2_list = [x for x in clean1_list if len(x) > 5]
    clean2_list = [x for x in clean2_list if str(x) != '  ']
    return clean2_list

def EscRes(fileName, title, imageName):
    df = pd.read_csv(fileName, delimiter=',', usecols=['Assignment group'])
    Helpdesk = 'IS 4th Source Helpdesk'
    Esc_ct = df[(df['Assignment group'] != Helpdesk)].count()
    Res_ct = df[(df['Assignment group'] == Helpdesk)].count()
    plt.bar('Esc', Esc_ct, label='Escalated', color='steelblue')
    plt.bar('Res', Res_ct, label='Resolved', color='darkorange')
    plt.title(title)
    plt.legend()
    plt.savefig("File_{}.png".format(imageName))
    plt.show()

def DailyTicketCount(fileName, imageName):
    df = pd.read_csv(fileName, delimiter=',')
    #convert dataframe to date & count frequency of days
    df = pd.to_datetime(df['Opened'])
    df1 = df.groupby(df.dt.date).count()
    print(df1)
    df2 = df1.plot(kind="bar", figsize=(14,6), fontsize=9)
    df2.set_alpha(0.8)
    df2.set_title("", fontsize=9)
    df2.set_xlabel("\nSunday - Saturday", fontsize=9)
    #df.set_xticklabels(['Sunday', 'Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday'], rotation=45)
    for p in df2.patches:
        df2.annotate ("%.0f" % p.get_height(), (p.get_x() + p.get_width() / 2., p.get_height()), ha='center', va='center', xytext=(0, 10), textcoords='offset points')
    plt.subplots_adjust(left=None, right=None, bottom=.2, wspace=0, hspace=0)
    plt.savefig("File_{}.png".format(imageName), bbox_inches='tight')
    plt.show()

def AsgmtGroup(fileName, imageName):
    df = pd.read_csv(fileName)
    ax = df['Assignment group'].value_counts().plot(kind='barh', figsize=(14,6), color="Orange", fontsize=10);
    ax.set_alpha(0.8)
    ax.set_title("", fontsize=10)
    ax.set_xlabel(" \nTickets", fontsize=10);
    #ax.set_xticks([0, 5, 10, 15, 20, 25, 30])
    #find the plt.patches values and append to list
    totals = []
    for i in ax.patches:
        totals.append(i.get_width())
    #set individual bar lables using totals list
    total = sum(totals)
    #set individual bar labels using totals list
    for i in ax.patches:
        # get_width pulls left or right; get_y pushes up or down
        ax.text(i.get_width()+.3, i.get_y()+.38, \
                str(round((i.get_width()/total)*100, 2))+'%', fontsize=10,color='dimgrey')
    #invert for largest on top 
    ax.invert_yaxis()
    #configure subplot parameters
    plt.subplots_adjust(left=None, bottom=.25, right=None, top=None, wspace=None, hspace=None)
    plt.savefig("File_{}.png".format(imageName), bbox_inches='tight')
    #show graph
    plt.show()

def CallStat1(fileName, graphTitle, imageName):
    df = pd.read_csv(fileName, delimiter=',')
    #find total calls
    totalCalls = find_total_calls(df)
    #turn dataframe to list & clean list
    cleaned_list = clean_list(df, totalCalls)
    #configure bar chart
    newDF = pd.DataFrame(list(cleaned_list))
    xlabels = (['Calls Placed on Hold', 'Transfers', 'Voicemails', 'Abandoned Calls', 'Total Calls']) 
    ax = (newDF).plot(kind='bar', color='blueviolet', figsize=(10,5), fontsize=9)
    ax.set_title(graphTitle, fontsize=10)
    ax.set_xticklabels(xlabels, fontsize=10, rotation=45)
    ax.get_legend().remove()
    rects = ax.patches
    #configure date formatted labels
    for rect, label in zip(rects, cleaned_list):
        height = rect.get_height()
        ax.text(rect.get_x() + rect.get_width() / 2, height + .5, label, ha='center', va='bottom')
    plt.yticks([])
    plt.subplots_adjust(left=.01, bottom=.4, right=.99, top=None, wspace=None, hspace=None)
    plt.savefig("File_{}.png".format(imageName))
    plt.show()

def CallStat2(Iteration, fileName, graphTitle, CSxlabels, imageName):
    df = pd.read_csv(fileName, delimiter=',')
    #find total calls
    totalCalls = find_total_calls(df)
    #turn dataframe to list & clean list
    cleaned_list = clean_list_averages(df)
    #slice strings from list, separate slices to new lists, & convert values to int
    hrs_list = [x[0:2] for x in cleaned_list]
    mins_list = [x[2:4] for x in cleaned_list]
    secs_list = [x[4:6] for x in cleaned_list]
    hrs_list = list(map(int, hrs_list))
    mins_list = list(map(int, mins_list))
    secs_list = list(map(int, secs_list))
    #calculate total times and place values in to new lists
    hrsToSecs_list = []
    minsToSecs_list = []
    totalTimeAsSecs_list = []
    times_list = []
    for x in range(len(hrs_list)):
        hrsToSecs_list.append(hrs_list[x]*3600)
    for x in range(len(mins_list)):
        minsToSecs_list.append(mins_list[x]*60)
    for x in range(Iteration):
        totalTimeAsSecs_list.append(hrsToSecs_list[x] + minsToSecs_list[x] + secs_list[x])
    for x in range(Iteration):
        times_list.append(str(datetime.timedelta(seconds=totalTimeAsSecs_list[x])))
    #configure bar chart
    newDF = pd.DataFrame(list(totalTimeAsSecs_list))
    xlabels = CSxlabels
    ax = (newDF).plot(kind='bar', color='lightseagreen', figsize=(10,5), fontsize=9)
    ax.set_title(graphTitle, fontsize=10)
    ax.set_xticklabels(xlabels, fontsize=10, rotation=45)
    ax.get_legend().remove()
    rects = ax.patches
    #configure date formatted labels
    for rect, label in zip(rects, times_list):
        height = rect.get_height()
        ax.text(rect.get_x() + rect.get_width() / 2, height + 5, label, ha='center', va='bottom')
    plt.yticks([])
    plt.subplots_adjust(left=.01, bottom=.4, right=.99, top=None, wspace=None, hspace=None)
    plt.savefig("File_{}.png".format(imageName))
    plt.show()

def SJM_TickReports():
    df = pd.read_csv('SJM.csv',delimiter=',',usecols=['Number', 'Short description', 'Assignment group', 'Resolution notes'])
    #identify esc/res tickets & place in to dataframes
    Helpdesk = 'IS 4th Source Helpdesk'
    Esc = df[(df['Assignment group'] != Helpdesk)]
    Esc = Esc.sort_values(by=['Assignment group'])
    Res = df[(df['Assignment group'] == Helpdesk)]
    #convert esc/res dataframes to HTML
    Esc.to_html('SJHD_Escalated.html')
    Res.to_html('SJHD_Resolved.html')
    #open newly created html tables
    webbrowser.open_new_tab('SJHD_Escalated.html')
    webbrowser.open_new_tab('SJHD_Resolved.html')

def main():
    
    #convert to CSV
    convert_toCSV()

    #create EscRes graph
    EscRes('SJM.csv',
        'St. Jude Milli\nEscalated & Resolved Graph',
        'SJM_1')

    #create DailyTicketCount graph
    DailyTicketCount('SJM.csv', 'SJM_2')

    #create AsgmtGroup graph
    AsgmtGroup('SJM.csv', 'SJM_3')

    #create CallStat1 graph
    CallStat1('SJ Milli - Previous Week.csv',
        'St. Jude Milli Support\nCall Statistics\n',
        'SJM_4')

    #create CallStat2 graph
    CallStat2(7,'SJ Milli - Previous Week.csv',
        'St. Jude Milli Support\nCall Statistics (Averages)\n',
        ['Answer Speed', 'After Call Work Time', 'Call  Time', 'Hold Time', 'Queue Wait Time', 'Ring Time', 'Talk Time'],
        'SJM_5')

    #create TicketReports
    SJM_TickReports()

    def createDoc_SJM():
        #create document
        document = Document()
        document.add_heading('AFH Milli Support', 0)
        p = document.add_paragraph('The following email details St. Jude activity for the week of.')
        document.add_picture('File_SJM_1.png', width=Inches(2.5), height=Inches(2))
        document.add_heading('Daily Ticket Breakdown: ', level=3)
        document.add_picture('File_SJM_2.png', width=Inches(7.25), height=Inches(3))
        document.add_heading('Assignment Group Breakdown: ', level=3)
        document.add_picture('File_SJM_3.png', width=Inches(7.25), height=Inches(2.75))
        document.add_heading('Escalated/Resolved:', level=3)
        document.add_heading('Call Statistics:', level=3)
        document.add_picture('File_SJM_4.png', width=Inches(6.5), height=Inches(3))
        document.add_picture('File_SJM_5.png', width=Inches(6.5), height=Inches(3))
        #save document
        document.save('SJM.docx')
        
    createDoc_SJM()

main()
