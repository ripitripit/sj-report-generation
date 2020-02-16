from docx import Document
from docx.shared import Inches

def createDoc_SJHD():
    #create document with Python
    document = Document()
    document.add_heading('AFH Helpdesk Support', 0)
    p = document.add_paragraph('The following email details St. Jude activity for the week of.')
    document.add_picture('File_SJHD_1.png', width=Inches(2.5), height=Inches(2))
    document.add_heading('Daily Ticket Breakdown: ', level=3)
    document.add_picture('File_SJHD_2.png', width=Inches(7.25), height=Inches(3))
    document.add_heading('Assignment Group Breakdown: ', level=3)
    document.add_picture('File_SJHD_3.png', width=Inches(7.25), height=Inches(2.75))
    document.add_heading('Escalated/Resolved:', level=3)
    document.add_heading('Call Statistics:', level=3)
    document.add_picture('File_SJHD_5.png', width=Inches(6.5), height=Inches(3))
    document.add_picture('File_SJHD_6.png', width=Inches(6.5), height=Inches(3))
    document.add_heading('Tickets by Categories:', level=3)
    document.add_picture('File_SJHD_7.png', width=Inches(4.1), height=Inches(3))
    document.add_picture('File_SJHD_8.png', width=Inches(4.1), height=Inches(3))
    #save document
    document.save('SJHD.docx')

def createDoc_SJM():
    #create document with Python
    document = Document()
    document.add_heading('AFH Helpdesk Support', 0)
    p = document.add_paragraph('The following email details St. Jude activity for the week of.')
    document.add_picture('File_SJM_1.png', width=Inches(2.5), height=Inches(2))
    document.add_heading('Daily Ticket Breakdown: ', level=3)
    document.add_picture('File_SJM_2.png', width=Inches(7.25), height=Inches(3))
    document.add_heading('Assignment Group Breakdown: ', level=3)
    document.add_picture('File_SJM_3.png', width=Inches(7.25), height=Inches(2.75))
    document.add_heading('Escalated/Resolved:', level=3)
    document.add_heading('Call Statistics:', level=3)
    document.add_picture('File_SJM_4.png', width=Inches(6.5), height=Inches(3))
    document.add_picture('File_SJM_5.png', width=Inches(6.5), height=Inches(3))
    #save document
    document.save('SJM.docx')

createDoc_SJHD()
createDoc_SJM()
